import docx
import re
from typing import List

class DOCXProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the DOCX processor with chunk settings
        
        Args:
            chunk_size (int): Maximum size of each text chunk
            chunk_overlap (int): Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process(self, file_path: str) -> List[str]:
        """
        Process a DOCX file and extract text chunks
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            List[str]: List of text chunks extracted from the DOCX file
        """
        try:
            # Extract text from DOCX
            text = self._extract_text(file_path)
            
            # Clean the text
            text = self._clean_text(text)
            
            # Split text into chunks
            chunks = self._chunk_text(text)
            
            return chunks
        
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text from the DOCX file
        """
        document = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text
        
        Args:
            text (str): Raw text extracted from DOCX
            
        Returns:
            str: Cleaned text
        """
        # Replace multiple newlines with single newline
        text = re.sub(r'\n+', '\n', text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove unnecessary Unicode characters
        text = text.encode('ascii', 'ignore').decode('ascii')
        
        return text.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text (str): Text to be chunked
            
        Returns:
            List[str]: List of text chunks
        """
        chunks = []
        
        if len(text) <= self.chunk_size:
            chunks.append(text)
            return chunks
        
        # Split text into chunks with overlap
        start = 0
        while start < len(text):
            # Find end of chunk
            end = start + self.chunk_size
            
            # Adjust end to not break mid-sentence if possible
            if end < len(text):
                # Try to end at a period, question mark, or exclamation point
                punctuation = max(text.rfind('. ', start, end), 
                                text.rfind('? ', start, end),
                                text.rfind('! ', start, end))
                
                if punctuation != -1:
                    end = punctuation + 1
            
            # Get chunk and add to list
            chunk = text[start:min(end, len(text))].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Move start position for next chunk
            start = end - self.chunk_overlap
            
            # Handle case where overlap might create a start position that's invalid
            if start < 0 or start >= len(text):
                break
        
        return chunks
